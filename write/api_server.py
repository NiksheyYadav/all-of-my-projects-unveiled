"""
Flask API Server for Research Writer
Provides REST API endpoints for the web interface
"""

from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from research_writer import ResearchWriter
import os
import json
import re
from pathlib import Path
from io import BytesIO

# Document conversion imports
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown2
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER

app = Flask(__name__, static_folder='.')
CORS(app)

# Initialize research writer
writer = ResearchWriter()

@app.route('/')
def index():
    """Serve the main HTML page"""
    return send_from_directory('.', 'index.html')

@app.route('/styles.css')
def serve_css():
    """Serve CSS file"""
    return send_from_directory('.', 'styles.css')

@app.route('/app.js')
def serve_js():
    """Serve JavaScript file"""
    return send_from_directory('.', 'app.js')

@app.route('/api/research', methods=['POST'])
def conduct_research():
    """Start a new research project"""
    try:
        data = request.json
        topic = data.get('topic')
        depth = data.get('depth', 'comprehensive')
        
        if not topic:
            return jsonify({'error': 'Topic is required'}), 400
        
        # Update config if settings provided
        if 'settings' in data:
            writer.config['research_settings'].update(data['settings'])
        
        # Conduct research
        result = writer.conduct_research(topic, depth, data.get('settings'))
        
        return jsonify({
            'success': True,
            'result': {
                'research_id': result['research_plan']['research_id'],
                'topic': topic,
                'sources_collected': result['sources_collected'],
                'sources_used': result['sources_used'],
                'document_path': result['document_path'],
                'plan': result['research_plan']
            }
        })
        
    except Exception as e:
        print(f"Error in conduct_research: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/document/<research_id>', methods=['GET'])
def get_document(research_id):
    """Get the generated research document"""
    try:
        # Find the document
        doc_path = Path('research_output') / f'research_{research_id}.md'
        
        if not doc_path.exists():
            return jsonify({'error': 'Document not found'}), 404
        
        with open(doc_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return jsonify({
            'success': True,
            'content': content,
            'research_id': research_id
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/results/<research_id>', methods=['GET'])
def get_results(research_id):
    """Get complete research results"""
    try:
        result_path = Path('research_output') / f'result_{research_id}.json'
        
        if not result_path.exists():
            return jsonify({'error': 'Results not found'}), 404
        
        with open(result_path, 'r') as f:
            results = json.load(f)
        
        return jsonify({
            'success': True,
            'results': results
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/config', methods=['GET', 'POST'])
def config_management():
    """Get or update configuration"""
    if request.method == 'GET':
        return jsonify({
            'success': True,
            'config': writer.config
        })
    else:
        try:
            data = request.json
            writer.config.update(data)
            
            # Save updated config
            with open('research_config.json', 'w') as f:
                json.dump(writer.config, f, indent=2)
            
            return jsonify({
                'success': True,
                'message': 'Configuration updated'
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500

@app.route('/api/history', methods=['GET'])
def get_history():
    """Get list of all research projects"""
    try:
        output_dir = Path('research_output')
        results = []
        
        for file in output_dir.glob('result_*.json'):
            try:
                with open(file, 'r') as f:
                    data = json.load(f)
                    results.append({
                        'research_id': data['research_plan']['research_id'],
                        'topic': data['research_plan']['topic'],
                        'created_at': data['research_plan']['created_at'],
                        'sources_used': data['sources_used']
                    })
            except Exception as e:
                print(f"Error reading {file}: {e}")
        
        # Sort by creation date (newest first)
        results.sort(key=lambda x: x['created_at'], reverse=True)
        
        return jsonify({
            'success': True,
            'history': results
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/history/<research_id>', methods=['DELETE'])
def delete_history(research_id):
    """Delete a research project"""
    try:
        # Files to delete
        files = [
            Path('research_output') / f'research_{research_id}.md',
            Path('research_output') / f'result_{research_id}.json',
            Path('research_output') / f'plan_{research_id}.json'
        ]
        
        deleted_count = 0
        for file in files:
            if file.exists():
                os.remove(file)
                deleted_count += 1
                
        return jsonify({
            'success': True,
            'message': f'Deleted {deleted_count} files'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/export/<format_type>/<research_id>', methods=['GET'])
def export_document(format_type, research_id):
    """Export research document in specified format (word or pdf)"""
    try:
        # Find the markdown document
        doc_path = Path('research_output') / f'research_{research_id}.md'
        
        if not doc_path.exists():
            return jsonify({'error': 'Document not found'}), 404
        
        with open(doc_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Get topic for filename
        result_path = Path('research_output') / f'result_{research_id}.json'
        topic = 'research'
        if result_path.exists():
            with open(result_path, 'r') as f:
                result_data = json.load(f)
                topic = result_data.get('research_plan', {}).get('topic', 'research')
        
        # Sanitize filename
        safe_filename = re.sub(r'[^a-zA-Z0-9_\s-]', '', topic)
        safe_filename = re.sub(r'\s+', '_', safe_filename)[:50]
        
        if format_type.lower() == 'word':
            # Convert to Word
            buffer = markdown_to_word(markdown_content, topic)
            return send_file(
                buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'{safe_filename}.docx'
            )
        elif format_type.lower() == 'pdf':
            # Convert to PDF
            buffer = markdown_to_pdf(markdown_content, topic)
            return send_file(
                buffer,
                mimetype='application/pdf',
                as_attachment=True,
                download_name=f'{safe_filename}.pdf'
            )
        else:
            return jsonify({'error': 'Invalid format. Use "word" or "pdf"'}), 400
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def markdown_to_word(markdown_content, title):
    """Convert markdown content to Word document"""
    doc = Document()
    
    # Add title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Parse markdown and convert to Word structure
    lines = markdown_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 4)
        # Bullet lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            content = re.sub(r'^\d+\.\s', '', line)
            doc.add_paragraph(content, style='List Number')
        # Regular paragraph
        else:
            # Handle bold and italic
            para = doc.add_paragraph()
            text = line
            # Simple bold/italic handling
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Remove markdown bold
            text = re.sub(r'\*(.+?)\*', r'\1', text)  # Remove markdown italic
            para.add_run(text)
        
        i += 1
    
    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def markdown_to_pdf(markdown_content, title):
    """Convert markdown content to PDF document"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=18)
    
    # Container for PDF elements
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    heading1_style = ParagraphStyle(
        'CustomHeading1',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=10,
        fontName='Helvetica-Bold'
    )
    
    # Add title
    elements.append(Paragraph(title, title_style))
    elements.append(Spacer(1, 0.2*inch))
    
    # Parse markdown and convert to PDF
    lines = markdown_content.split('\n')
    for line in lines:
        line = line.rstrip()
        
        if not line:
            elements.append(Spacer(1, 0.1*inch))
            continue
        
        # Clean up markdown syntax for PDF
        text = line
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)  # Bold
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)  # Italic
        
        # Headers
        if line.startswith('# '):
            elements.append(Paragraph(text[2:], heading1_style))
        elif line.startswith('## '):
            elements.append(Paragraph(text[3:], heading2_style))
        elif line.startswith('### '):
            elements.append(Paragraph(text[4:], styles['Heading3']))
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            elements.append(Paragraph('• ' + text[2:], styles['Normal']))
        elif re.match(r'^\d+\.\s', line):
            elements.append(Paragraph(text, styles['Normal']))
        # Regular paragraph
        else:
            elements.append(Paragraph(text, styles['Normal']))
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer

if __name__ == '__main__':
    print("🚀 Starting Research Writer API Server")
    print("📡 Server running at http://localhost:8000")
    print("📝 Open http://localhost:8000 in your browser")
    app.run(debug=True, port=8000)
